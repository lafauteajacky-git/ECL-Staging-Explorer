"""Committee summary generation for ECL Staging Explorer V0.6."""

from __future__ import annotations

from io import BytesIO

import pandas as pd


REQUIRED_SECTIONS = [
    "## 1. Resume executif",
    "## 2. Vue d'ensemble du portefeuille",
    "## 3. Resultats de staging",
    "## 4. Resultats ECL",
    "## 5. Analyse des scenarios macroeconomiques",
    "## 6. Analyse des overlays",
    "## 7. Qualite des donnees et points d'attention",
    "## 8. Conclusion et recommandations",
    "## 9. Client Discussion Points",
]


def format_amount(value: float) -> str:
    """Format amounts for the markdown committee note."""
    return f"{value:,.0f} EUR".replace(",", " ")


def generate_committee_summary(
    run_id: str,
    metrics: dict,
    scenario_metrics: dict,
    overlay_metrics: dict,
    ecl_by_stage: pd.DataFrame,
    ecl_by_product: pd.DataFrame,
    ecl_by_sector: pd.DataFrame,
    stage_counts: pd.DataFrame,
    scenario_parameters: pd.DataFrame,
    scenario_summary: pd.DataFrame,
    overlay_summary: pd.DataFrame,
    data_quality_summary: pd.DataFrame,
    review_count: int,
    top_contributors: pd.DataFrame,
    insights: list[str],
    business_consistency_summary: dict | None = None,
    client_discussion_points: list[str] | None = None,
    demo_profile: str | None = None,
) -> str:
    """Generate a sober markdown committee summary from calculated results."""
    key_message = insights[0] if insights else "Les resultats doivent etre interpretes dans le cadre simplifie du MVP."
    top_products = _top_labels(ecl_by_product, "product_type", "ecl")
    top_sectors = _top_labels(ecl_by_sector, "sector", "ecl")
    stage_lines = "\n".join(
        f"- {row['stage']}: {int(row['exposure_count'])} expositions, EAD {format_amount(row['ead'])}, ECL {format_amount(row['ecl'])}"
        for _, row in ecl_by_stage.iterrows()
    )
    scenario_lines = "\n".join(
        f"- {row['scenario']}: poids {row['weight']:.0%}, ECL {format_amount(row['ecl'])}"
        for _, row in scenario_summary.iterrows()
    )
    overlay_lines = "\n".join(
        f"- {row['overlay_name']}: {format_amount(row['overlay_amount'])} ({row['justification']})"
        for _, row in overlay_summary.iterrows()
        if row["overlay_amount"] > 0
    ) or "- Aucun overlay avec impact non nul."
    dq_lines = "\n".join(
        f"- {row['check_code']}: {int(row['issue_count'])} anomalie(s)"
        for _, row in data_quality_summary.iterrows()
    ) or "- Aucune anomalie detectee."
    top_lines = "\n".join(
        f"- {row['loan_id']}: {format_amount(row['ecl'])}, stage {row['stage']}"
        for _, row in top_contributors.head(5).iterrows()
    )
    business_consistency_summary = business_consistency_summary or {}
    client_discussion_points = client_discussion_points or []
    discussion_lines = "\n".join(f"- {point}" for point in client_discussion_points) or "- Non disponible."
    consistency_score = business_consistency_summary.get("business_consistency_score", 1.0)
    alert_count = business_consistency_summary.get("business_alert_count", 0)
    critical_count = business_consistency_summary.get("business_critical_alert_count", 0)

    return f"""# Note de synthese - Comite provisionnement

Run ID: {run_id}
Version de demonstration: {demo_profile or "Non specifie"}

## 1. Resume executif

- EAD totale: {format_amount(metrics['total_ead'])}
- ECL finale apres overlays: {format_amount(overlay_metrics['ecl_after_overlay'])}
- Taux de couverture modele: {metrics['coverage_ratio']:.2%}
- Variation liee aux scenarios: {format_amount(scenario_metrics['weighted_impact_amount'])} ({scenario_metrics['weighted_impact_pct']:.2%})
- Variation liee aux overlays: {format_amount(overlay_metrics['total_overlay_amount'])} ({overlay_metrics['overlay_variation_pct']:.2%})
- Score de coherence metier: {consistency_score:.1%}
- Message cle: {key_message}

## 2. Vue d'ensemble du portefeuille

- Nombre d'expositions: {metrics['exposure_count']}
- Principaux produits contributeurs: {top_products}
- Principaux secteurs contributeurs: {top_sectors}

{stage_lines}

## 3. Resultats de staging

- Stage 1 / Stage 2 / Stage 3: voir repartition ci-dessus.
- Regles les plus frequemment declenchees: voir onglet Staging Results.
- Cas necessitant revue: {review_count}

## 4. Resultats ECL

- ECL avant scenarios: {format_amount(metrics['total_ecl'])}
- ECL ponderee apres scenarios: {format_amount(scenario_metrics['ecl_weighted'])}
- ECL avant overlays: {format_amount(overlay_metrics['ecl_before_overlay'])}
- ECL finale apres overlays: {format_amount(overlay_metrics['ecl_after_overlay'])}
- Taux de couverture global modele: {metrics['coverage_ratio']:.2%}

Top contributeurs:
{top_lines}

## 5. Analyse des scenarios macroeconomiques

{scenario_lines}

- Impact downside vs baseline: {format_amount(scenario_metrics['downside_impact_amount'])} ({scenario_metrics['downside_impact_pct']:.2%})
- Impact ECL ponderee vs baseline: {format_amount(scenario_metrics['weighted_impact_amount'])} ({scenario_metrics['weighted_impact_pct']:.2%})

## 6. Analyse des overlays

- Montant total des overlays: {format_amount(overlay_metrics['total_overlay_amount'])}
- Impact en pourcentage: {overlay_metrics['overlay_variation_pct']:.2%}
- Overlay le plus significatif: {overlay_metrics['top_overlay_contributor']}

{overlay_lines}

## 7. Qualite des donnees et points d'attention

- Nombre total d'anomalies: {metrics['data_quality_issue_count']}
- Expositions a revoir: {review_count}
- Alertes de coherence metier: {alert_count} dont {critical_count} critique(s)
- Impact potentiel: les anomalies critiques peuvent affecter la fiabilite du calcul et doivent etre analysees avant usage production.

{dq_lines}

## 8. Conclusion et recommandations

- Valider les principaux contributeurs ECL avec les equipes Risques et Finance.
- Revoir les expositions marquees `review_required`.
- Documenter formellement les hypotheses de scenarios et d'overlays.
- Completer les controles de coherence avant une industrialisation.
- Garder en memoire que ce MVP est un demonstrateur sur donnees synthetiques.

## 9. Client Discussion Points

{discussion_lines}
"""


def _top_labels(df: pd.DataFrame, label_col: str, value_col: str) -> str:
    if df.empty:
        return "non disponible"
    return ", ".join(df.sort_values(value_col, ascending=False).head(3)[label_col].astype(str).tolist())


def build_docx_bytes(markdown_text: str) -> bytes:
    """Build a simple DOCX file from markdown-like text."""
    from docx import Document

    document = Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
